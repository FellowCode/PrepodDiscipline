import os

from Disciplines.models import Fakultet
from Prepods.models import Prepod
from utils.shortcuts import get_group_nagruzki, annotate_pochasovka, annotate_group_nagruzki, get_shtat_rasp
from docx import Document
from shutil import copyfile


def word_shtat_rasp(request, fakultet_id=None, _all=False):
    if fakultet_id == '':
        fakultet_id = None
    fakultet = Fakultet.objects.get_or_none(id=fakultet_id)
    prepods = get_shtat_rasp(request, fakultet, _all)
    doc = Document('static/files/Shtat_raspisanie.docx')
    print(doc.tables)

    for i, row in enumerate(prepods['rows']):
        doc.tables[0].cell(i+14, 0).text = str(i+1)
        for j, col in enumerate(row):
            doc.tables[0].cell(i + 14, j+1).text = str(col)


    doc.tables[0].cell(34, 8).text = str(prepods['sums']['n_stavka_sum'])
    doc.tables[0].cell(34, 9).text = f"{prepods['sums']['n_p_stavka_sum']}\n({prepods['sums']['n_p_ch_stavka_sum']})"
    doc.tables[0].cell(34, 10).text = str(prepods['sums']['v_n_stavka_sum'])
    doc.tables[0].cell(34, 11).text = f"{prepods['sums']['v_n_p_stavka_sum']}\n({prepods['sums']['v_n_p_ch_stavka_sum']})"
    doc.tables[0].cell(35, 8).text = str(prepods['sums']['n_stavka_sum'] + prepods['sums']['n_p_stavka_sum'])
    doc.tables[0].cell(35, 10).text = str(prepods['sums']['v_n_stavka_sum'] + prepods['sums']['v_n_p_stavka_sum'])
    doc.tables[0].cell(36, 8).text = str(prepods['sums']['n_stavka_sum'] + prepods['sums']['n_p_stavka_sum'] + prepods['sums']['v_n_stavka_sum'] + prepods['sums']['v_n_p_stavka_sum'])


    for j in range(18-len(prepods['rows'])):
        row = doc.tables[0].rows[len(prepods['rows'])+14]
        remove_row(doc.tables[0], row)


    if not os.path.exists(f'tmp/word/{request.user.id}'):
        os.makedirs(f'tmp/word/{request.user.id}')

    if request.user.prepod.count() > 0:
        kafedra = request.user.prepod.first().kafedra.name
    else:
        kafedra = 'Все кафедры'

    if fakultet:
        filename = f'{fakultet.name} - штатное расписание.doc'
    else:
        filename = f'{kafedra} - штатное расписание.doc'
    path = f'tmp/word/{request.user.id}/{filename}'

    doc.save(path)

    return path


def remove_row(table, row):
    tbl = table._tbl
    tr = row._tr
    tbl.remove(tr)